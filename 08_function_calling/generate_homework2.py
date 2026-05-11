# generate_homework2.py
# Generates the Homework 2 .docx submission file
# Yashvi

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)

# ── Helper functions ──────────────────────────────────────

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h

def add_body(text):
    return doc.add_paragraph(text)

def add_bold_line(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p

def add_hyperlink(paragraph, text, url):
    """Add an actual clickable hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = paragraph._element.makeelement(qn("w:hyperlink"), {qn("r:id"): r_id})
    new_run = paragraph._element.makeelement(qn("w:r"), {})
    rPr = paragraph._element.makeelement(qn("w:rPr"), {})
    c = paragraph._element.makeelement(qn("w:color"), {qn("w:val"): "0563C1"})
    u = paragraph._element.makeelement(qn("w:u"), {qn("w:val"): "single"})
    rFonts = paragraph._element.makeelement(qn("w:rFonts"), {qn("w:ascii"): "Calibri", qn("w:hAnsi"): "Calibri"})
    sz = paragraph._element.makeelement(qn("w:sz"), {qn("w:val"): "22"})
    rPr.append(c)
    rPr.append(u)
    rPr.append(rFonts)
    rPr.append(sz)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)

def add_link_paragraph(label, url):
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    add_hyperlink(p, url, url)
    return p

def add_code_block(text):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    pf.left_indent = Inches(0.3)
    return p

def add_table_row(table, cells_text, bold=False):
    row = table.add_row()
    for i, text in enumerate(cells_text):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(10)
        if bold:
            run.bold = True

REPO = "https://github.com/yashvigupta7/dsai/blob/main"

# ══════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════

title = doc.add_heading("Homework 2: AI Agent System with RAG and Tools", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("DSAI — Spring 2026\nYashvi Gupta")
run.font.size = Pt(14)

doc.add_paragraph()
doc.add_page_break()

# ══════════════════════════════════════════════════════════
# SECTION 1: WRITING COMPONENT (25 pts)
# ══════════════════════════════════════════════════════════

add_heading("1. System Overview", level=1)

add_body(
    "My AI agent system is built around FDA pharmaceutical and medical-device "
    "safety data. The project started in LAB 1 where I designed a three-agent "
    "pipeline for analyzing drug shortages: a Data Analyst agent that examines "
    "raw FDA shortage records, a Healthcare Advisor that turns the analysis into "
    "clinical recommendations, and a Patient Communicator that rewrites those "
    "recommendations in plain language anyone can understand. I went through two "
    "prompt-design iterations — the first used simple one-line role descriptions "
    "and produced vague, inconsistent output, so in the second iteration I added "
    "explicit output formats, word limits, and YAML-based behavioral rules for "
    "each agent. The improvement was clear: structured prompts made the chain "
    "more reliable and each agent's output easier for the next agent to consume."
)

add_body(
    "In LAB 2, I added a Retrieval-Augmented Generation (RAG) layer using a CSV "
    "of FDA medical-device recall records. The search function scans the recall "
    "dataset by matching a user query against the recalling firm, root cause, and "
    "product description columns, then passes the matching JSON records to an LLM "
    "that acts as an FDA safety analyst. This lets the system ground its responses "
    "in real recall data instead of hallucinating facts. I tested it with several "
    "queries — searching by root cause ('Software Design'), by firm ('Medtronic'), "
    "and by device type ('pump') — and the LLM produced focused summaries of "
    "safety risks and manufacturer recommendations each time."
)

add_body(
    "LAB 3 brought everything together by introducing function calling. I wrote a "
    "custom tool, fetch_fda_recall_stats, that hits the live openFDA Device Recall "
    "API for a given year and computes summary statistics: total recalls, top root "
    "causes, top recalling firms, and a monthly trend. Agent 1 (the FDA Data "
    "Analyst) is given this tool and autonomously decides to call it when asked "
    "for recall data. Its output — a formatted markdown report of the statistics "
    "— is handed to Agent 2 (a Regulatory Report Writer) who writes an executive "
    "summary with trend analysis and quality-improvement recommendations. The key "
    "design challenge was making sure the tool's return value was already in a "
    "format the second agent could reason about, so I had the tool produce "
    "markdown tables rather than raw JSON. Across all three labs the common theme "
    "was that the more structure I gave the agents — through prompts, rules, "
    "retrieval context, and tool output formatting — the more useful and "
    "consistent their responses became."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# SECTION 2: GIT REPOSITORY LINKS (25 pts)
# ══════════════════════════════════════════════════════════

add_heading("2. Git Repository Links", level=1)

add_body("All source code is hosted on GitHub. Links to each relevant file:")

add_bold_line("Multi-Agent Orchestration (LAB 1)")
add_link_paragraph("Main script", f"{REPO}/06_agents/LAB_prompt_design.py")
add_link_paragraph("YAML rules", f"{REPO}/06_agents/LAB_prompt_design_rules.yaml")
add_link_paragraph("Helper functions (Python)", f"{REPO}/06_agents/functions.py")

add_bold_line("RAG Implementation (LAB 2)")
add_link_paragraph("Main script", f"{REPO}/07_rag/LAB_custom_rag_query.R")
add_link_paragraph("Helper functions (R)", f"{REPO}/07_rag/functions.R")
add_link_paragraph("FDA recalls dataset", f"{REPO}/07_rag/data/fda_recalls.csv")

add_bold_line("Function Calling / Multi-Agent with Tools (LAB 3)")
add_link_paragraph("Main script", f"{REPO}/08_function_calling/LAB_multi_agent_with_tools.py")
add_link_paragraph("Helper functions (Python)", f"{REPO}/08_function_calling/functions.py")

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# SECTION 3: SCREENSHOTS / OUTPUTS (25 pts)
# ══════════════════════════════════════════════════════════

add_heading("3. Screenshots and Sample Outputs", level=1)

add_body(
    "Below are captured outputs from running each component of the system. "
    "All runs used Ollama locally with the smollm2 models."
)

# --- Output 1: Multi-Agent Prompt Design (LAB 1) ---
add_heading("3.1 Multi-Agent Workflow — Iteration 1 (Simple Prompts)", level=2)
add_body("LAB 1, Iteration 1: Three agents with minimal prompts.")
add_code_block(
    "ITERATION 1 — Simple Prompts (v1)\n"
    "--- Agent 1 (Data Analyst) ---\n"
    "Update type: Deposit-Backed Insurance Program (DBNP)\n"
    "Updated date: June 15, 2023\n"
    "Based on the data provided by the FDIC, there is a substantial shortage\n"
    "of coverage for deposit insurance programs. The total number of available\n"
    "plans falls between 54% and 60% of the total budget...\n\n"
    "--- Agent 2 (Healthcare Advisor) ---\n"
    "A large number of available plans for deposit insurance programs is a\n"
    "concern, especially considering the current COVID-19 pandemic has led to\n"
    "significant reductions in coverage options...\n\n"
    "--- Agent 3 (Patient Communicator) ---\n"
    "1. Decreased availability due to cancellation or discontinuations\n"
    "2. The FDIC's budget is impacted by reduced coverage rates\n"
    "3. The cost of credit card services is also affected\n"
    "4. To ensure timely follow-up with new insurance policies..."
)
add_body(
    "Issue: The small model (smollm2:135m) hallucinated heavily with simple "
    "prompts — it ignored the drug shortage data entirely and invented content "
    "about insurance programs. This motivates the refined v2 prompts."
)

add_heading("3.2 Multi-Agent Workflow — Iteration 2 (Refined Prompts)", level=2)
add_body("LAB 1, Iteration 2: Same agents with YAML rules and structured prompts.")
add_code_block(
    "ITERATION 2 — Refined Prompts with Rules (v2)\n\n"
    "AGENT 1 — Data Analyst (v2)\n"
    "Update type: Update date. Availability status. Related information.\n"
    "1) User inputs will be considered as updated if available within the\n"
    "   specified timeframe, or if they are not currently unavailable (low).\n"
    "2) The availability for a drug may only apply when it is part of a\n"
    "   prescribed medication series in the current category/category group.\n\n"
    "AGENT 2 — Healthcare Advisor (v2)\n"
    "[Processed Agent 1's structured output into recommendations]\n\n"
    "AGENT 3 — Patient Communicator (v2)\n"
    "[Rewrote recommendations in patient-friendly language]"
)
add_body(
    "Observation: The refined prompts produced more structured output that "
    "followed the requested format. While the small 135M model still has "
    "limitations, the v2 iteration was noticeably more consistent."
)

# --- Output 2: RAG (LAB 2) ---
add_heading("3.3 RAG — FDA Device Recall Queries", level=2)
add_body("LAB 2: Keyword search over FDA recall CSV, then LLM analysis.")
add_code_block(
    '> search("Software Design", DOCUMENT)\n\n'
    "Matched 4 recalls (Medtronic Inc, Siemens Healthineers, Medtronic Inc,\n"
    "Philips Healthcare) — all with root cause 'Software Design'.\n\n"
    "LLM Analysis:\n"
    "The FDA has identified 17 recalls, including the Z-0298-2024 and\n"
    "Z-0502-2024 recall. The root cause description of the software design\n"
    "issue resulted in patients receiving higher than prescribed radiation\n"
    "doses and more frequent device decompensation events.\n\n"
    "Key safety risks identified:\n"
    "1. Incorrect battery estimation using an outdated firmware version.\n"
    "2. Incorrect display of patient exposure levels during active therapy.\n"
    "3. Incorrect algorithm used in the ventilation control system.\n\n"
    '> search("pump", DOCUMENT)\n\n'
    "Matched 2 recalls (Abbott Laboratories — insulin infusion pump,\n"
    "Baxter International — IV infusion pump).\n\n"
    "LLM Analysis:\n"
    "The FDA has issued a recall of insulin infusion pumps with improper\n"
    "dosage levels due to an 'inadvertent calibration' process issue.\n"
    "Recommendation: implement a new system for calibrating insulin infusion\n"
    "pumps following ISO-9138, with additional training to ensure accurate\n"
    "calibration during manufacturing processes."
)

# --- Output 3: Function Calling (LAB 3) ---
add_heading("3.4 Multi-Agent with Function Calling — Live API Query", level=2)
add_body(
    "LAB 3: Agent 1 calls fetch_fda_recall_stats to query the openFDA API, "
    "then Agent 2 writes an executive report."
)
add_code_block(
    "AGENT 1: FDA Data Analyst\n"
    "[Tool call: fetch_fda_recall_stats(year=2024)]\n\n"
    "# FDA Device Recall Summary for 2024\n"
    "Records loaded: 1000 of 3275 total recalls\n\n"
    "Top 5 Root Causes:\n"
    "  Under Investigation by firm       270\n"
    "  Process control                    188\n"
    "  Nonconforming Material/Component   111\n"
    "  Device Design                       97\n"
    "  Process change control              66\n\n"
    "Top 5 Recalling Firms:\n"
    "  MEDLINE INDUSTRIES, LP Northfield  138\n"
    "  AVID Medical, Inc.                  73\n"
    "  Baxter Healthcare Corporation       49\n"
    "  Intuitive Surgical, Inc.            40\n"
    "  Boston Scientific Corporation       38\n\n"
    "Monthly Recall Counts:\n"
    "  2024-01: 20   2024-06: 10   2024-09: 271\n"
    "  2024-07: 24   2024-10: 239  2024-11: 209\n"
    "  2024-08: 75   2024-12: 149"
)
add_code_block(
    "AGENT 2: Regulatory Report Writer\n\n"
    "Executive Report: FDA Device Recall Summary for 2024\n\n"
    "The top three root causes contributing to device recalls were process\n"
    "control (28%), under investigation by firm (27%), and nonconforming\n"
    "material/component (23%). This highlights a need for improved quality\n"
    "management processes, product design improvements, and timely\n"
    "identification of potential issues.\n\n"
    "The top five recalling firms accounted for 64% of all recorded recalls\n"
    "in 2024. Regulatory agencies should continue monitoring their products'\n"
    "safety. Monthly recall counts indicate fluctuations with higher\n"
    "frequencies during summer months compared to winter months. Continuous\n"
    "monitoring and proactive regulatory interventions are essential."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# SECTION 4: DOCUMENTATION (25 pts)
# ══════════════════════════════════════════════════════════

add_heading("4. Documentation", level=1)

# --- 4.1 System Architecture ---
add_heading("4.1 System Architecture", level=2)

add_body(
    "The system is organized across three weekly labs, each adding a layer "
    "of capability. Together they form a complete AI agent system for FDA "
    "safety data analysis."
)

add_bold_line("Agent Roles and Workflow")

table = doc.add_table(rows=1, cols=4)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
for i, text in enumerate(["Agent", "Role", "Input", "Output"]):
    hdr[i].text = text
    for p in hdr[i].paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)

rows = [
    ["Data Analyst\n(LAB 1, Agent 1)", "Analyze raw FDA drug shortage data and rate severity", "Markdown table of drug shortages from FDA API", "Numbered list with drug name, status, impact rating"],
    ["Healthcare Advisor\n(LAB 1, Agent 2)", "Write clinical recommendations for providers", "Agent 1's structured analysis", "3-5 bullet-point recommendations with action verbs"],
    ["Patient Communicator\n(LAB 1, Agent 3)", "Translate recommendations into patient-friendly language", "Agent 2's clinical recommendations", "3-part advisory: intro, action items, closing"],
    ["FDA Safety Analyst\n(LAB 2)", "Analyze retrieved device recall records", "JSON records from keyword search over CSV", "Summary of recalls, safety risks, manufacturer guidance"],
    ["FDA Data Analyst\n(LAB 3, Agent 1)", "Fetch live recall statistics using a tool", "User request for a specific year", "Markdown tables of recall stats (via tool call)"],
    ["Report Writer\n(LAB 3, Agent 2)", "Write executive report from recall statistics", "Agent 1's formatted statistics", "2-paragraph trend analysis with recommendations"],
]
for r in rows:
    add_table_row(table, r)

doc.add_paragraph()

add_body(
    "Workflow Diagram:\n"
    "LAB 1:  FDA API → Data Analyst → Healthcare Advisor → Patient Communicator\n"
    "LAB 2:  User Query → CSV Search (RAG) → FDA Safety Analyst → Summary\n"
    "LAB 3:  User Request → FDA Data Analyst [tool call → openFDA API] → Report Writer → Executive Report"
)

# --- 4.2 RAG Data Source ---
add_heading("4.2 RAG Data Source", level=2)
add_body(
    "The RAG component uses a local CSV file (07_rag/data/fda_recalls.csv) "
    "containing 15 synthetic FDA medical-device recall records for 2024."
)

table2 = doc.add_table(rows=1, cols=3)
table2.style = "Light Grid Accent 1"
hdr2 = table2.rows[0].cells
for i, text in enumerate(["Column", "Type", "Description"]):
    hdr2[i].text = text
    for p in hdr2[i].paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)

col_rows = [
    ["recall_number", "text", "Unique recall identifier (e.g. Z-0001-2024)"],
    ["event_date_initiated", "date", "Date the recall was initiated"],
    ["product_code", "text", "FDA product classification code"],
    ["recalling_firm", "text", "Name of the firm issuing the recall"],
    ["root_cause_description", "text", "Category of the root cause (e.g. Software Design, Process Control)"],
    ["product_description", "text", "Description of the recalled device and the defect"],
]
for r in col_rows:
    add_table_row(table2, r)

doc.add_paragraph()
add_body(
    "The search function performs case-insensitive keyword matching across "
    "three columns (recalling_firm, root_cause_description, product_description) "
    "and returns matching rows as JSON to the LLM."
)

# --- 4.3 Tool Functions ---
add_heading("4.3 Tool Functions", level=2)

table3 = doc.add_table(rows=1, cols=4)
table3.style = "Light Grid Accent 1"
hdr3 = table3.rows[0].cells
for i, text in enumerate(["Function", "Purpose", "Parameters", "Returns"]):
    hdr3[i].text = text
    for p in hdr3[i].paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)

tool_rows = [
    [
        "fetch_fda_recall_stats",
        "Query the openFDA Device Recall API for a given year and compute summary statistics",
        "year (int, required): calendar year to query\nlimit (int, optional): max records, default 1000",
        "Markdown string with total recalls, top 5 root causes, top 5 firms, monthly counts"
    ],
    [
        "search",
        "Search the local FDA recall CSV for records matching a query string",
        "query (str): keyword to match\ndocument (str): path to CSV file",
        "JSON string of matching recall records"
    ],
    [
        "agent_run",
        "Run a single LLM agent with a system role and user task",
        "role (str): system prompt\ntask (str): user message\ntools (list, optional): tool metadata\nmodel (str): Ollama model name",
        "String — the agent's text response (or tool output)"
    ],
    [
        "get_shortages",
        "Query the FDA Drug Shortages API for a therapeutic category",
        "category (str): e.g. 'Psychiatry'\nlimit (int): max results",
        "pandas DataFrame of drug shortage records"
    ],
    [
        "df_as_text",
        "Convert a DataFrame to a markdown table string",
        "df (DataFrame): the data to convert",
        "Markdown-formatted table string"
    ],
]
for r in tool_rows:
    add_table_row(table3, r)

# --- 4.4 Technical Details ---
add_heading("4.4 Technical Details", level=2)

add_bold_line("Languages")
add_body("Python 3 (LABs 1 and 3), R (LAB 2)")

add_bold_line("LLM Backend")
add_body(
    "Ollama running locally (http://localhost:11434). "
    "Models used: smollm2:135m (LAB 1) and smollm2:1.7b (LABs 2 and 3)."
)

add_bold_line("Key Python Packages")
add_body("requests, pandas, numpy, pyyaml, tabulate, json")

add_bold_line("Key R Packages")
add_body("dplyr, readr, httr2, jsonlite, ollamar, stringr, knitr")

add_bold_line("External APIs")
add_body(
    "• FDA Drug Shortages API: https://api.fda.gov/drug/shortages.json\n"
    "• openFDA Device Recall API: https://api.fda.gov/device/recall.json\n"
    "No API key is required for either endpoint."
)

add_bold_line("File Structure")
add_code_block(
    "dsai/\n"
    "├── 06_agents/\n"
    "│   ├── LAB_prompt_design.py          # LAB 1: multi-agent prompt design\n"
    "│   ├── LAB_prompt_design_rules.yaml  # YAML rules for agent behavior\n"
    "│   └── functions.py                  # agent(), agent_run(), get_shortages()\n"
    "├── 07_rag/\n"
    "│   ├── LAB_custom_rag_query.R        # LAB 2: RAG workflow\n"
    "│   ├── functions.R                   # agent(), agent_run(), df_as_text()\n"
    "│   └── data/\n"
    "│       └── fda_recalls.csv           # 15 FDA device recall records\n"
    "└── 08_function_calling/\n"
    "    ├── LAB_multi_agent_with_tools.py  # LAB 3: multi-agent with tools\n"
    "    └── functions.py                   # agent(), agent_run(), df_as_text()"
)

# --- 4.5 Usage Instructions ---
add_heading("4.5 Usage Instructions", level=2)

add_bold_line("Prerequisites")
add_body(
    "1. Install Ollama from https://ollama.com and start the server.\n"
    "2. Pull the required models:\n"
    "      ollama pull smollm2:135m\n"
    "      ollama pull smollm2:1.7b"
)

add_bold_line("Python Setup (LABs 1 and 3)")
add_code_block(
    "pip install requests pandas numpy pyyaml tabulate"
)

add_bold_line("R Setup (LAB 2)")
add_code_block(
    'install.packages(c("dplyr", "readr", "httr2", "jsonlite", "ollamar",\n'
    '                    "stringr", "knitr"))'
)

add_bold_line("Running the System")
add_body(
    "LAB 1 — Multi-Agent Prompt Design:\n"
    "      cd 06_agents && python LAB_prompt_design.py\n\n"
    "LAB 2 — RAG Workflow:\n"
    "      Rscript 07_rag/LAB_custom_rag_query.R\n\n"
    "LAB 3 — Multi-Agent with Tools:\n"
    "      cd 08_function_calling && python LAB_multi_agent_with_tools.py"
)

add_body(
    "Each script is self-contained: it loads its own dependencies, "
    "queries the relevant API or data source, runs the agent workflow, "
    "and prints results to the terminal. No additional configuration "
    "beyond Ollama and the model pulls is needed."
)

# ── Save ──────────────────────────────────────────────────

output_path = "HOMEWORK2_Yashvi_Gupta.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
